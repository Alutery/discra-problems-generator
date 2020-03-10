#!/usr/bin/python
# -*- coding: UTF-8 -*-

'''
Example:

In the card game of bridge, you are dealt a hand of 5 cards from the standard 36-card deck. 
How many different hands are there where the player has 5 cards in one suit and 2 cards in each of the other suits? 
(like 5H+2D+2S+2C).

 '''
import sys
from math import factorial
import argparse

# required: sudo pip3 install python-docx 
from docx import Document

# сочетаниями из n объектов по k
def C(k, n):
    return factorial(n) // factorial(n-k) // factorial(k)

class Generator:

    def __init__(self):
        self.doc = None
        self.min_n = None
        self.max_n = None

    # для обработки аргументов командной строки
    def createParser (self):
        parser = argparse.ArgumentParser()
        parser.add_argument ('--min', nargs='?', default=0, type=int, help='Минимальный возможный ответ (по умолчанию 0)')
        parser.add_argument ('--max', nargs='?', default=25000000, type=int, help='Максимальный возможный ответ (по умолчанию 25000000)')
        parser.add_argument ('-f', '--file', default='mit_cards_1', help='Наименование файла документа docx (! без .docx, по умолчанию mit_cards_1)')
    
        return parser

    # проверка принадлежности ответа заданному диапазону
    def check(self, n):
        return True if (n <= self.max_n and n >= self.min_n) else False

    # получть задания по заданному количеству карт
    def get_tasks(self, x, y, z=1, type_task=1):

        if x > 9 or y > 9 or x < 0 or y < 0:
            print(f'Invalid Arguments! ({x}, {y})')
            return

        text = ''
        if type_task == 1:
            text += f'In the card game of bridge, you are dealt a hand of 5 cards from the standard 36-card deck. \
    How many different hands are there where the player has {x} and {y} card(s) in some suits, respectively, and {z} card(s) in each of the other suits? \
    (like {x}H+{y}D+{z}S+{z}C).\n\n'
            culc = f'C(1, 4) * C(1, 3) * ( C({x}, 9) * C({y}, 9) * C({z}, 9)^2 )'
            text += culc + '\n\nAnswer: '
            answer = eval(f'C(1, 4) * C(1, 3) * ( C({x}, 9) * C({y}, 9) * C({z}, 9) ** 2 )')
            if not self.check(answer):
                return None
            text += str(answer)

        else:
            text += f'In the card game of bridge, you are dealt a hand of 5 cards from the standard 36-card deck. \
    How many different hands are there where the player has {x} card(s) in one suit and {y} card(s) in each of the other suits? \
    (like {x}H+{y}D+{y}S+{y}C).\n\n'
            culc = f'C(1, 4) * ( C({x}, 9) * C({y}, 9)^3 )'
            text += culc + '\n\nAnswer : '
            answer = eval(f'C(1, 4) * ( C({x}, 9) * C({y}, 9) ** 3 )')
            if not self.check(answer):
                return None
            text += str(answer)

        return text + '\n'

    # генерация заданий
    def generate_tasks(self):
        count = 1
        self.doc.add_heading('Часть 1.', 0)
        self.doc.add_paragraph()

        for i in range(1, 10):
            for j in range(1, 10):
                if i == j:
                    continue
                task = self.get_tasks(i, j, type_task=2)
                if task:
                    self.doc.add_paragraph(str(count) + '. ' + task)
                    count += 1

        self.doc.add_page_break()
        self.doc.add_heading('Часть 2.', 0)
        self.doc.add_paragraph()

        count = 1
        for i in range(1, 10):
            for j in range(1, 10):
                for k in range(1, 10):
                    if i == j or i == k or j == k:
                        continue
                    task = self.get_tasks(i, j, k, type_task=1)
                    if task:
                        self.doc.add_paragraph(str(count) + '. ' + task)
                        count += 1

    # создание документа
    def generate(self):
        parser = self.createParser()
        namespace = parser.parse_args(sys.argv[1:])

        self.min_n = namespace.min
        self.max_n = namespace.max

        self.doc = Document()
        self.generate_tasks()
        self.doc.save(namespace.file + '.docx')


if __name__ == '__main__':
    Generator().generate()