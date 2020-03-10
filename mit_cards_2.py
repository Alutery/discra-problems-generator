#!/usr/bin/python
# -*- coding: UTF-8 -*-

'''
Example:

Two identical complete decks of cards, each contains 36 cards, have been mixed together. 
A hand of a cards is picked uniformly at random from amongst all subsets of exactly a cards.
Identical cards are cards with the same suit and value. 
For example, the hand  (Q♥, 5♠, 6♠, 8♣, Q♥) has identical cards.) 

What is the probability that the hand has exactly 2 pairs of identical cards?

Comment:
Probability is the ratio of the number of considered hands-configurations to the number of all hands-configurations.
We distinguish cards from different decks!

Write only 3 digits after the decimal point. 
(If the answer is 0.0449999… then write “044”, neither “0.044” nor “045”).

 '''
import sys
from math import factorial
import argparse

# required: sudo pip3 install python-docx 
from docx import Document


# сочетаниями из n объектов по k
def C(k, n):
    return factorial(n) // factorial(n-k) // factorial(k)


class Generator:

    def __init__(self):
        self.doc = None

    # для обработки аргументов командной строки
    def createParser (self):
        parser = argparse.ArgumentParser()
        parser.add_argument ('-d', '--desk', nargs='*', default=36, type=int, help='Число карт в одной колоде (по умолчанию 36)')
        parser.add_argument ('-f', '--file', default='mit_cards_2', help='Наименование файла документа docx (! без .docx, по умолчанию mit_cards_2)')
    
        return parser

    # получть задания по заданному количеству карт
    def get_task(self, n_pairs=1, hand=5, desk=36, type_task=1):

        if n_pairs * 2 > hand:
            # print('Invalid Arguments!')
            return

        text = f'''Two identical complete decks of cards, each contains {desk} cards, have been mixed together. 
    A hand of {hand} cards is picked uniformly at random from amongst all subsets of exactly {hand} cards.
    Identical cards are cards with the same suit and value. 
    For example, the hand  (Q♥, 5♠, 6♠, 8♣, Q♥) has identical cards.)\n\n'''

        text += f'What is the probability that the hand has exactly {n_pairs} pair(s) of identical cards?\n\n'

        all_cases = f'C({hand}, {desk} * 2)'
        culc = f'( C({n_pairs}, {desk}) * C({hand - 2 * n_pairs}, {desk - n_pairs}) * ( C(1, 2)^{hand - 2 * n_pairs} ) ) / ' + all_cases
        text += culc + '\n\nAnswer: '
        answer = eval(culc.replace('^', ' ** ', 1))

        number_dec = str(answer-int(answer))[2:5]
        text += number_dec + ' (' + str(answer) + ')'

        return text + '\n'

    # генерация заданий по заданной колоде
    def generate_desk(self, desk):
        count = 1
        self.doc.add_heading(f'Desk of {desk}.', 0)
        self.doc.add_paragraph()

        for hand in range(5, 8):
            for n_pairs in range(1, 4):
                task = self.get_task(n_pairs, hand, desk)
                if task:
                    self.doc.add_paragraph(str(count) + '. ' + task)
                    count += 1

        self.doc.add_page_break()


    def generate(self):
        parser = self.createParser()
        namespace = parser.parse_args(sys.argv[1:])

        self.doc = Document()

        desks = namespace.desk

        if isinstance(desks, int):
            desks = [desks]

        for desk in desks:
            self.generate_desk(desk)
        self.doc.save(namespace.file + '.docx')


if __name__ == '__main__':
    Generator().generate()